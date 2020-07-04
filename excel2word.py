"""
create time: 20200704
author: mba1398
task: use excel and stastic word template to generate a lot of word files. 
"""


from docx import Document
import xlrd


def text_chenge(headline, data):
    # 用来替换word段落中的关键字内容，关键字都是excel表格的标题行
    myparagraphs = document.paragraphs
    for paragraph in myparagraphs:
        for run in paragraph.runs:
            run_text = run.text.replace(headline, data)
            run.text = run_text

    # 用来替换word表格中的关键字内容，关键字都是excel表格的标题行
    mytables = document.tables
    for table in mytables:
        for row in table.rows:
            for cell in row.cells:
                cell_text = cell.text.replace(headline, data)
                cell.text = cell_text


xlsx = xlrd.open_workbook(r'D:\pycharm\learning\autowork\document\报告数据.xlsx')
table = xlsx.sheet_by_index(0)

for table_row in range(1, table.nrows):
    document = Document(r'D:\pycharm\learning\autowork\document\报告模板.docx')
    for table_col in range(0, table.ncols):
        text_chenge(str(table.cell(0, table_col).value), str(table.cell(table_row, table_col).value))
        # 将excel表格中的内容替换掉标题行，因为标题行即为报告模板中的关键字

    document.save(f'{str(table.cell(table_row, 0).value)} eSRVCC切换成功率低优化报告.docx')
    print("%s eSRVCC切换成功率低优化报告成功生成!" % str(table.cell_value(table_row, 0)))
    
